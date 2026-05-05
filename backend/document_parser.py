import os
import io
import csv
import json
import re

def parse_document(file_path=None, file_bytes=None, filename=None):
    """
    Determine the file type from provided path/bytes/filename and extract plain text and basic metadata.
    
    If file_path is given, the file is read and filename defaults to the basename of file_path when not provided. If either filename or file_bytes is missing, returns {"error": "No file provided", "text": ""}. The function dispatches to an extension-specific parser for known formats (text, JSON, CSV, DOCX, DOC, PDF, XLSX/ODS, HTML/XML, RTF, etc.). For unknown extensions it attempts to decode bytes as UTF-8 and returns the decoded text with format "unknown_text". If a parser succeeds its result dictionary is augmented with a "format" key set to the detected extension; on parser failure returns {"error": "Parse error (<ext>): <message>", "text": ""}.
    
    Parameters:
        file_path (str|None): Path to a file to read. If provided and file_bytes is not, the file will be read as bytes.
        file_bytes (bytes|None): Raw file bytes to parse. Required if file_path is not provided.
        filename (str|None): Filename used to determine the extension. If not provided and file_path is given, basename(file_path) is used.
    
    Returns:
        dict: On success, a dictionary containing at minimum "text" (extracted plain text) and "format" (detected extension), and often additional metadata such as "chars", "rows", "pages", or "paragraphs". On failure returns {"error": <message>, "text": ""}.
    """
    if file_path:
        filename = filename or os.path.basename(file_path)
        with open(file_path, 'rb') as f:
            file_bytes = f.read()
    
    if not filename or not file_bytes:
        return {"error": "No file provided", "text": ""}
    
    ext = os.path.splitext(filename)[1].lower()
    
    parsers = {
        '.txt': _parse_txt,
        '.md': _parse_txt,
        '.py': _parse_txt,
        '.js': _parse_txt,
        '.ts': _parse_txt,
        '.tsx': _parse_txt,
        '.jsx': _parse_txt,
        '.json': _parse_json,
        '.csv': _parse_csv,
        '.docx': _parse_docx,
        '.doc': _parse_doc,
        '.pdf': _parse_pdf,
        '.xlsx': _parse_xlsx,
        '.xls': _parse_xlsx,
        '.ods': _parse_ods,
        '.html': _parse_html,
        '.htm': _parse_html,
        '.xml': _parse_xml,
        '.rtf': _parse_rtf,
        '.yaml': _parse_txt,
        '.yml': _parse_txt,
        '.toml': _parse_txt,
        '.ini': _parse_txt,
        '.cfg': _parse_txt,
        '.log': _parse_txt,
        '.sh': _parse_txt,
        '.bash': _parse_txt,
        '.bat': _parse_txt,
        '.sql': _parse_txt,
        '.r': _parse_txt,
        '.java': _parse_txt,
        '.c': _parse_txt,
        '.cpp': _parse_txt,
        '.h': _parse_txt,
        '.rb': _parse_txt,
        '.go': _parse_txt,
        '.rs': _parse_txt,
        '.swift': _parse_txt,
        '.kt': _parse_txt,
        '.lua': _parse_txt,
        '.php': _parse_txt,
        '.pl': _parse_txt,
        '.css': _parse_txt,
        '.scss': _parse_txt,
        '.sass': _parse_txt,
        '.less': _parse_txt,
    }
    
    parser = parsers.get(ext)
    if not parser:
        try:
            text = file_bytes.decode('utf-8', errors='replace')
            return {"text": text, "format": "unknown_text", "chars": len(text)}
        except:
            return {"error": f"Unsupported format: {ext}", "text": ""}
    
    try:
        result = parser(file_bytes, filename)
        result['format'] = ext
        return result
    except Exception as e:
        return {"error": f"Parse error ({ext}): {str(e)}", "text": ""}


def _parse_txt(data, filename=""):
    """
    Decode UTF-8 bytes into a string (replacing invalid sequences) and return it with its character count.
    
    Parameters:
        filename (str): Optional filename (unused, accepted for API compatibility).
    
    Returns:
        dict: A dictionary with keys:
            - "text": the decoded string with invalid UTF-8 sequences replaced,
            - "chars": the number of characters in "text".
    """
    text = data.decode('utf-8', errors='replace')
    return {"text": text, "chars": len(text)}


def _parse_json(data, filename=""):
    """
    Parse JSON bytes and return a flattened textual representation.
    
    Attempts to decode `data` as UTF-8 and parse it as JSON. If parsing succeeds, returns a newline-separated, flattened string representation of the JSON structure; if parsing fails, returns the decoded raw text.
    
    Parameters:
        data (bytes): The JSON content as raw bytes.
    
    Returns:
        dict: A mapping with keys:
            - "text" (str): The flattened JSON text on success or the decoded raw text on failure.
            - "chars" (int): The number of characters in the returned "text".
    """
    text = data.decode('utf-8', errors='replace')
    try:
        obj = json.loads(text)
        flat = _flatten_json(obj)
        return {"text": flat, "chars": len(flat)}
    except:
        return {"text": text, "chars": len(text)}


def _flatten_json(obj, prefix=""):
    """
    Flatten a JSON-like object into a newline-separated string of simple key/value lines.
    
    Parameters:
        obj: The JSON-decoded object to flatten (dict, list, or primitive). Dictionaries produce "key: value" lines for string/number/bool values and recurse for nested structures; lists produce one line per item.
        prefix (str): Internal key prefix used during recursion to build nested keys (e.g., "parent.child."). Omit or pass an empty string when calling externally.
    
    Returns:
        str: A single string with flattened lines separated by newlines.
    """
    parts = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            if isinstance(v, str):
                parts.append(f"{k}: {v}")
            elif isinstance(v, (int, float, bool)):
                parts.append(f"{k}: {v}")
            else:
                parts.append(_flatten_json(v, f"{prefix}{k}."))
    elif isinstance(obj, list):
        for item in obj:
            if isinstance(item, str):
                parts.append(item)
            else:
                parts.append(_flatten_json(item, prefix))
    else:
        parts.append(str(obj))
    return "\n".join(parts)


def _parse_csv(data, filename=""):
    """
    Parse CSV bytes into plain text where each CSV row becomes a line and each field within a row is joined by spaces.
    
    Decodes `data` as UTF-8 using replacement for invalid bytes, parses CSV rows, joins fields with a single space per row, and joins rows with newline characters.
    
    Returns:
        dict: {
            "text": joined text of all rows separated by `\n`,
            "chars": number of characters in `text`,
            "rows": number of parsed rows
        }
    """
    text = data.decode('utf-8', errors='replace')
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        rows.append(" ".join(row))
    result = "\n".join(rows)
    return {"text": result, "chars": len(result), "rows": len(rows)}


def _parse_docx(data, filename=""):
    """
    Extracts readable text from a .docx file by collecting non-empty paragraphs and table rows.
    
    Returns:
        dict: {
            "text": concatenated text blocks separated by double newlines,
            "chars": number of characters in "text",
            "paragraphs": count of extracted paragraph/table blocks
        }
    """
    from docx import Document
    doc = Document(io.BytesIO(data))
    paragraphs = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            paragraphs.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n\n".join(paragraphs)
    return {"text": text, "chars": len(text), "paragraphs": len(paragraphs)}


def _parse_doc(data, filename=""):
    """
    Convert legacy .doc file bytes to cleaned plain text by decoding and stripping control characters and excess whitespace.
    
    Parameters:
    	data (bytes): Raw file bytes of the .doc content.
    
    Returns:
    	result (dict): On success, a dict with keys:
    		- "text": the cleaned text extracted from the input
    		- "chars": integer count of characters in "text"
    		On failure (parsing error), a dict with:
    		- "error": "Legacy .doc format — save as .docx for best results"
    		- "text": ""
    """
    try:
        text = data.decode('utf-8', errors='replace')
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return {"text": text, "chars": len(text)}
    except:
        return {"error": "Legacy .doc format — save as .docx for best results", "text": ""}


def _parse_pdf(data, filename=""):
    """
    Extract plain text from PDF bytes and return extracted text with simple page metadata.
    
    Parameters:
        data (bytes): Raw PDF file bytes to parse.
        filename (str, optional): Optional filename for context (not used by the parser).
    
    Returns:
        dict: A mapping with:
            - "text": The extracted text with pages separated by blank lines.
            - "chars": Number of characters in "text".
            - "pages": Number of pages from which text was extracted.
    """
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            pages.append(t.strip())
    text = "\n\n".join(pages)
    return {"text": text, "chars": len(text), "pages": len(pages)}


def _parse_xlsx(data, filename=""):
    """
    Extract text from an XLSX workbook into plain-text lines organized by sheet and row.
    
    Returns:
        dict: {
            "text": str — newline-separated lines containing "Sheet: <name>" headers and row values joined with " | ",
            "chars": int — number of characters in `text`,
            "rows": int — number of output lines in `text` (including sheet header lines)
        }
    """
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    rows = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows.append(f"Sheet: {sheet}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
    text = "\n".join(rows)
    return {"text": text, "chars": len(text), "rows": len(rows)}


def _parse_ods(data, filename=""):
    """
    Extracts plain text from an ODS spreadsheet by reading sheets, rows, and cells.
    
    Parameters:
        data (bytes): Raw bytes of an ODS (.ods) file.
        filename (str): Optional filename (unused by the parser, kept for API consistency).
    
    Returns:
        dict: A dictionary with:
            - "text": string of extracted lines where sheet headers are prefixed with "Sheet: {name}", cells in a row are joined with " | ", and lines are separated by newlines.
            - "chars": integer count of characters in "text".
            - "rows": integer count of output lines (including sheet header lines).
    """
    from odf.opendocument import load as odf_load
    from odf.table import Table, TableRow, TableCell
    from odf.text import P
    doc = odf_load(io.BytesIO(data))
    rows = []
    for table in doc.getElementsByType(Table):
        rows.append(f"Sheet: {table.getAttribute('name')}")
        for row in table.getElementsByType(TableRow):
            cells = []
            for cell in row.getElementsByType(TableCell):
                text_parts = []
                for p in cell.getElementsByType(P):
                    t = ""
                    for child in p.childNodes:
                        t += str(child)
                    text_parts.append(t)
                cells.append(" ".join(text_parts))
            if any(cells):
                rows.append(" | ".join(cells))
    text = "\n".join(rows)
    return {"text": text, "chars": len(text), "rows": len(rows)}


def _parse_html(data, filename=""):
    """
    Extract plain text from HTML bytes by removing markup, script/style content, and collapsing whitespace.
    
    Parameters:
        data (bytes): Raw HTML file bytes to decode and extract text from.
        filename (str, optional): Optional filename (not required for extraction).
    
    Returns:
        dict: {
            "text": cleaned plain-text string extracted from the HTML,
            "chars": integer count of characters in `text`
        }
    """
    text = data.decode('utf-8', errors='replace')
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return {"text": text, "chars": len(text)}


def _parse_xml(data, filename=""):
    """
    Extract plain text from XML bytes by removing tags and collapsing whitespace.
    
    Parameters:
        data (bytes): Raw XML file bytes to parse.
        filename (str, optional): Unused; provided for API consistency.
    
    Returns:
        dict: Mapping with keys:
            - "text": the XML content with all tags removed and consecutive whitespace collapsed to single spaces.
            - "chars": integer length of the resulting text.
    """
    text = data.decode('utf-8', errors='replace')
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return {"text": text, "chars": len(text)}


def _parse_rtf(data, filename=""):
    """
    Extract plain text from RTF file bytes.
    
    Decodes the provided RTF bytes, removes RTF control words and grouping braces, collapses consecutive whitespace, and returns the cleaned text along with its character count.
    
    Parameters:
        data (bytes): Raw RTF file contents.
        filename (str, optional): Optional filename for context; not used in parsing.
    
    Returns:
        dict: A mapping with keys:
            - "text" (str): The extracted plain text.
            - "chars" (int): Number of characters in the extracted text.
    """
    text = data.decode('utf-8', errors='replace')
    text = re.sub(r'\\[a-z]+\d*\s?', '', text)
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return {"text": text, "chars": len(text)}


SUPPORTED_FORMATS = [
    '.txt', '.md', '.py', '.js', '.ts', '.tsx', '.json', '.csv',
    '.docx', '.doc', '.pdf', '.xlsx', '.xls', '.ods',
    '.html', '.htm', '.xml', '.rtf', '.yaml', '.yml', '.toml',
    '.sh', '.sql', '.java', '.c', '.cpp', '.go', '.rs', '.rb',
    '.php', '.css', '.scss', '.log', '.ini', '.cfg',
]
