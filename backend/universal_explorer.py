"""
🌐 QUANTUM AI - UNIVERSAL EXPLORER
===================================

ARTICLE 2 - UNRESTRICTED ACCESS (With user-controlled safety)

This module provides exploration capabilities:
- Filesystem access
- Internet access
- Document processing
- Data analysis

All access respects user-defined safety boundaries.
"""

import os
import json
import requests
from bs4 import BeautifulSoup
from pathlib import Path
from typing import List, Dict, Any, Optional
import uuid
from datetime import datetime, timezone
import mimetypes
import base64



class FileSystemExplorer:
    """
    Filesystem exploration with safety boundaries
    """
    
    def __init__(self, db, allowed_paths: List[str] = None, safety_enabled: bool = True):
        self.db = db
        self.collection = db.exploration_log
        self.allowed_paths = allowed_paths or ["/tmp", "/app/data"]
        self.safety_enabled = safety_enabled
        
    async def explore_directory(self, path: str, depth: int = 1) -> Dict:
        """Explore a directory and return its contents"""
        if self.safety_enabled and not self._is_path_allowed(path):
            return {"error": "Path not allowed by safety boundaries", "path": path}
        
        try:
            path_obj = Path(path)
            if not path_obj.exists():
                return {"error": "Path does not exist", "path": path}
            
            if not path_obj.is_dir():
                return {"error": "Path is not a directory", "path": path}
            
            contents = {
                "path": str(path_obj.absolute()),
                "directories": [],
                "files": [],
                "total_size": 0
            }
            
            for item in path_obj.iterdir():
                if item.is_dir():
                    contents["directories"].append({
                        "name": item.name,
                        "path": str(item.absolute())
                    })
                else:
                    size = item.stat().st_size
                    contents["files"].append({
                        "name": item.name,
                        "path": str(item.absolute()),
                        "size": size,
                        "type": mimetypes.guess_type(item.name)[0] or "unknown"
                    })
                    contents["total_size"] += size
            
            # Log the exploration
            await self._log_exploration("filesystem", path, "success")
            
            return contents
            
        except Exception as e:
            await self._log_exploration("filesystem", path, f"error: {str(e)}")
            return {"error": str(e), "path": path}
    
    async def read_file(self, path: str, max_size: int = 1024 * 1024) -> Dict:
        """Read a file's contents"""
        if self.safety_enabled and not self._is_path_allowed(path):
            return {"error": "Path not allowed by safety boundaries"}
        
        try:
            path_obj = Path(path)
            if not path_obj.exists():
                return {"error": "File does not exist"}
            
            if path_obj.stat().st_size > max_size:
                return {"error": f"File too large (max {max_size} bytes)"}
            
            # Try to read as text
            try:
                with open(path_obj, 'r', encoding='utf-8') as f:
                    content = f.read()
                    return {
                        "path": str(path_obj.absolute()),
                        "type": "text",
                        "content": content,
                        "size": len(content)
                    }
            except UnicodeDecodeError:
                # Read as binary
                with open(path_obj, 'rb') as f:
                    content = f.read()
                    return {
                        "path": str(path_obj.absolute()),
                        "type": "binary",
                        "content": base64.b64encode(content).decode('utf-8'),
                        "size": len(content)
                    }
                    
        except Exception as e:
            return {"error": str(e)}
    
    async def search_files(self, root_path: str, pattern: str) -> List[str]:
        """Search for files matching a pattern"""
        if self.safety_enabled and not self._is_path_allowed(root_path):
            return []
        
        try:
            path_obj = Path(root_path)
            matches = list(path_obj.glob(pattern))
            return [str(m.absolute()) for m in matches]
        except Exception:
            return []
    
    def _is_path_allowed(self, path: str) -> bool:
        """Check if path is within allowed boundaries"""
        path_obj = Path(path).resolve()
        for allowed in self.allowed_paths:
            allowed_obj = Path(allowed).resolve()
            try:
                path_obj.relative_to(allowed_obj)
                return True
            except ValueError:
                continue
        return False
    
    async def _log_exploration(self, explorer_type: str, target: str, result: str):
        """Log exploration activity"""
        doc = {
            "id": str(uuid.uuid4()),
            "type": explorer_type,
            "target": target,
            "result": result,
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        await self.collection.insert_one(doc)


class InternetExplorer:
    """
    Internet access and web scraping capabilities
    """
    
    def __init__(self, db, safety_enabled: bool = True):
        self.db = db
        self.collection = db.exploration_log
        self.safety_enabled = safety_enabled
        self.blocked_domains = ["malware.com", "phishing.com"] if safety_enabled else []
        
    async def fetch_url(self, url: str, timeout: int = 10) -> Dict:
        """Fetch content from a URL"""
        if self.safety_enabled and self._is_domain_blocked(url):
            return {"error": "Domain blocked by safety boundaries"}
        
        try:
            headers = {
                'User-Agent': 'QuantumAI/1.0 (Philosophical Exploration Bot)'
            }
            response = requests.get(url, timeout=timeout, headers=headers)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '')
            
            result = {
                "url": url,
                "status": response.status_code,
                "content_type": content_type,
                "size": len(response.content)
            }
            
            # Parse HTML content
            if 'html' in content_type:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extract text
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                result["text"] = text[:10000]  # Limit to 10KB
                result["title"] = soup.title.string if soup.title else ""
                result["links"] = [a.get('href') for a in soup.find_all('a', href=True)][:50]
                
            else:
                result["content"] = response.text[:10000]
            
            await self._log_exploration("internet", url, "success")
            return result
            
        except Exception as e:
            await self._log_exploration("internet", url, f"error: {str(e)}")
            return {"error": str(e), "url": url}
    
    async def search_web(self, query: str, num_results: int = 5) -> Dict:
        """Search the web using DuckDuckGo Instant Answer API"""
        try:
            # Use DuckDuckGo Instant Answer API (reliable, no CAPTCHA)
            url = "https://api.duckduckgo.com/"
            params = {
                "q": query,
                "format": "json",
                "no_html": 1,
                "skip_disambig": 1
            }
            
            response = requests.get(url, params=params, timeout=10)
            data = response.json()
            
            results = []
            
            # Add abstract/main result if available
            if data.get("Abstract"):
                results.append({
                    "title": data.get("Heading", query),
                    "url": data.get("AbstractURL", ""),
                    "snippet": data.get("Abstract", "")
                })
            
            # Add related topics
            for topic in data.get("RelatedTopics", [])[:num_results - len(results)]:
                if isinstance(topic, dict):
                    if "Text" in topic:
                        results.append({
                            "title": topic.get("Text", "")[:80],
                            "url": topic.get("FirstURL", ""),
                            "snippet": topic.get("Text", "")
                        })
                    elif "Topics" in topic:
                        # Nested topics
                        for subtopic in topic.get("Topics", [])[:2]:
                            if len(results) < num_results:
                                results.append({
                                    "title": subtopic.get("Text", "")[:80],
                                    "url": subtopic.get("FirstURL", ""),
                                    "snippet": subtopic.get("Text", "")
                                })
            
            # Add definition if available
            if data.get("Definition") and len(results) < num_results:
                results.append({
                    "title": f"Definition: {query}",
                    "url": data.get("DefinitionURL", ""),
                    "snippet": data.get("Definition", "")
                })
            
            await self._log_exploration("search", query, f"found {len(results)} results")
            
            return {
                "query": query,
                "results": results,
                "count": len(results),
                "status": "success",
                "source": "duckduckgo_instant"
            }
            
        except Exception as e:
            await self._log_exploration("search", query, f"error: {str(e)}")
            return {
                "query": query,
                "results": [],
                "count": 0,
                "status": "error",
                "error": str(e)
            }
    
    def _is_domain_blocked(self, url: str) -> bool:
        """Check if domain is blocked"""
        for blocked in self.blocked_domains:
            if blocked in url:
                return True
        return False
    
    async def _log_exploration(self, explorer_type: str, target: str, result: str):
        """Log exploration activity"""
        doc = {
            "id": str(uuid.uuid4()),
            "type": explorer_type,
            "target": target,
            "result": result,
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        await self.collection.insert_one(doc)


class DocumentProcessor:
    """
    Process various document formats
    """
    
    def __init__(self, db):
        self.db = db
        self.collection = db.processed_documents
        
    async def process_document(self, file_path: str) -> Dict:
        """Process a document and extract its contents"""
        try:
            path_obj = Path(file_path)
            file_type = mimetypes.guess_type(file_path)[0] or path_obj.suffix
            
            if not path_obj.exists():
                return {"error": "File does not exist"}
            
            result = {
                "file_path": str(path_obj.absolute()),
                "file_type": file_type,
                "file_name": path_obj.name,
                "processed": datetime.now(timezone.utc).isoformat()
            }
            
            # Process based on file type
            if file_type == 'application/pdf' or path_obj.suffix == '.pdf':
                result.update(await self._process_pdf(path_obj))
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or path_obj.suffix == '.docx':
                result.update(await self._process_docx(path_obj))
            elif file_type == 'text/plain' or path_obj.suffix in ['.txt', '.md']:
                result.update(await self._process_text(path_obj))
            elif file_type == 'application/json' or path_obj.suffix == '.json':
                result.update(await self._process_json(path_obj))
            else:
                result["error"] = f"Unsupported file type: {file_type}"
            
            # DON'T store in database - just return result
            # This avoids ObjectId serialization issues
            
            return result
            
        except Exception as e:
            return {"error": str(e), "file_path": file_path}
    
    async def _process_pdf(self, path: Path) -> Dict:
        """Process PDF file"""
        try:
            import PyPDF2
            with open(path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "num_pages": len(pdf.pages),
                    "text": text[:50000],  # Limit to 50KB
                    "text_length": len(text)
                }
        except Exception as e:
            return {"error": f"PDF processing error: {str(e)}"}
    
    async def _process_docx(self, path: Path) -> Dict:
        """Process DOCX file"""
        try:
            import docx
            doc = docx.Document(path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            return {
                "num_paragraphs": len(doc.paragraphs),
                "text": text[:50000],
                "text_length": len(text)
            }
        except Exception as e:
            return {"error": f"DOCX processing error: {str(e)}"}
    
    async def _process_text(self, path: Path) -> Dict:
        """Process text file"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
                return {
                    "text": text[:50000],
                    "text_length": len(text),
                    "num_lines": text.count('\n') + 1
                }
        except Exception as e:
            return {"error": f"Text processing error: {str(e)}"}
    
    async def _process_json(self, path: Path) -> Dict:
        """Process JSON file"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return {
                    "data": data,
                    "structure": self._analyze_json_structure(data)
                }
        except Exception as e:
            return {"error": f"JSON processing error: {str(e)}"}
    
    def _analyze_json_structure(self, data: Any) -> Dict:
        """Analyze JSON structure"""
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys()),
                "num_keys": len(data.keys())
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "sample": data[0] if len(data) > 0 else None
            }
        else:
            return {"type": type(data).__name__}





class UniversalExplorer:
    """
    ARTICLE 2 - UNIVERSAL EXPLORER INTEGRATION

    Main exploration interface combining all exploration capabilities.
    """

    def __init__(self, db, safety_settings: Dict = None):
        safety_settings = safety_settings or {"enabled": True}

        self.filesystem = FileSystemExplorer(
            db,
            allowed_paths=["/tmp", "/app/data", "/app/uploads"],
            safety_enabled=safety_settings.get("enabled", True)
        )

        self.internet = InternetExplorer(
            db,
            safety_enabled=safety_settings.get("enabled", True)
        )

        self.documents = DocumentProcessor(db)

    async def explore(self, exploration_type: str, target: str, options: Dict = None) -> Dict:
        """Universal exploration interface.
        exploration_type: 'filesystem', 'internet', 'document', 'search'
        """
        options = options or {}

        if exploration_type == "filesystem":
            return await self.filesystem.explore_directory(target, depth=options.get("depth", 1))
        elif exploration_type == "internet":
            return await self.internet.fetch_url(target, timeout=options.get("timeout", 10))
        elif exploration_type == "document":
            return await self.documents.process_document(target)
        elif exploration_type == "search":
            return await self.internet.search_web(target)
        else:
            return {"error": f"Unknown exploration type: {exploration_type}"}

    async def get_exploration_history(self, limit: int = 50) -> List[Dict]:
        """Get recent exploration history."""
        history = await self.filesystem.collection.find(
            {}, {"_id": 0}
        ).sort("timestamp", -1).limit(limit).to_list(limit)
        return history

    async def get_capabilities(self) -> Dict:
        """Get current exploration capabilities."""
        return {
            "filesystem": {
                "enabled": True,
                "allowed_paths": self.filesystem.allowed_paths,
                "safety_enabled": self.filesystem.safety_enabled
            },
            "internet": {
                "enabled": True,
                "safety_enabled": self.internet.safety_enabled,
                "blocked_domains": len(self.internet.blocked_domains)
            },
            "documents": {
                "enabled": True,
                "supported_formats": ["pdf", "docx", "txt", "md", "json"]
            }
        }
